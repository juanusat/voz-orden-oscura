from docx import Document
from docx.shared import Pt
from pathlib import Path


def generate_docx(transcription, output_path, options=None):
    Path(Path(output_path).parent).mkdir(parents=True, exist_ok=True)
    doc = Document()
    title = options.get("title") if options else None
    if title:
        doc.add_heading(title, level=1)
    segments = None
    if isinstance(transcription, dict):
        segments = transcription.get("segments")
    else:
        segments = getattr(transcription, "segments", None)
    if isinstance(segments, list):
        for seg in segments:
            start = seg.get("start") if isinstance(seg, dict) else None
            end = seg.get("end") if isinstance(seg, dict) else None
            text = seg.get("text") if isinstance(seg, dict) else ""
            p = doc.add_paragraph()
            p.add_run(f"[{float(start):.2f}-{float(end):.2f}] ").bold = True
            r = p.add_run(text)
            r.font.size = Pt(11)
    else:
        text = getattr(transcription, "text", "") if not isinstance(transcription, dict) else transcription.get("text", "")
        doc.add_paragraph(text)
    doc.save(output_path)
    return output_path
